from pathlib import Path
from docx import Document

base = Path(r'C:\Users\meabh\OneDrive\Documents\Program\Codes\Project\Major Project by Python\My Self\Contract\App\mou_templates')
for p in sorted(base.glob('*.docx')):
    print(f'FILE: {p.name}')
    doc = Document(str(p))
    paras = [para.text.strip() for para in doc.paragraphs if para.text.strip()]
    for i, t in enumerate(paras[:80], start=1):
        print(f'P{i}: {t}')
    for ti, table in enumerate(doc.tables[:3], start=1):
        print(f'TABLE {ti}:')
        for row in table.rows[:8]:
            vals = [cell.text.strip() for cell in row.cells]
            if any(vals):
                print('  ', vals)
    print('---END FILE---')
    print()
