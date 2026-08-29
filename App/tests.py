from pathlib import Path
from types import SimpleNamespace

from django.contrib.auth import get_user_model
from django.test import TestCase
from django.urls import reverse
from docx import Document

from .models import FacilityInformation, MOUHistory
from .views import TEMPLATE_FILES, _build_replacements, replace_text_in_document


class DocumentReplacementTests(TestCase):
    def test_replaces_split_run_placeholder_and_table_text(self):
        document = Document()
        paragraph = document.add_paragraph()
        paragraph.add_run("{HOSPITAL_")
        paragraph.add_run("NAME}")
        table = document.add_table(rows=1, cols=1)
        table.cell(0, 0).text = "{AMBULANCE_NAME}"
        facility = SimpleNamespace(
            hospital_name="Central Hospital", hospital_address="1 Main Street",
            dry_cleaner_name="", dry_cleaner_address="", blood_bank_name="", blood_bank_address="",
            canteen_name="", canteen_address="", second_hospital_name="Referral Hospital",
            second_hospital_address="2 Main Street", ambulance_name="Rapid Ambulance",
            ambulance_address="3 Main Street", radio_lab_name="", radio_lab_address="",
            lab_name="", lab_address="",
        )

        replace_text_in_document(document, _build_replacements(facility))

        self.assertEqual(document.paragraphs[0].text, "Central Hospital")
        self.assertEqual(document.tables[0].cell(0, 0).text, "Rapid Ambulance")


class DashboardWorkflowTests(TestCase):
    def setUp(self):
        self.user = get_user_model().objects.create_user(username="manager", password="test-password")
        self.client.force_login(self.user)
        self.facility = FacilityInformation.objects.create(
            hospital_name="Central Hospital",
            hospital_address="1 Main Street",
            ambulance_name="Rapid Ambulance",
            ambulance_address="3 Main Street",
        )

    def tearDown(self):
        output_directory = Path("generated_mous")
        for path in output_directory.glob(f"*_MOU_{self.facility.id}.docx"):
            path.unlink(missing_ok=True)

    def test_dashboard_and_facility_pages_render(self):
        self.assertEqual(self.client.get(reverse("dashboard")).status_code, 200)
        self.assertEqual(self.client.get(reverse("facility_list")).status_code, 200)
        self.assertEqual(self.client.get(reverse("facility_detail", args=[self.facility.id])).status_code, 200)
        self.assertEqual(self.client.get(reverse("mou_documents")).status_code, 200)
        self.assertEqual(self.client.get(reverse("generate_mou_page")).status_code, 200)
        self.assertEqual(self.client.get(reverse("mou_history")).status_code, 200)

    def test_facility_edit_and_post_delete(self):
        response = self.client.post(reverse("facility_edit", args=[self.facility.id]), {
            "hospital_name": "Updated Hospital",
            "hospital_address": "2 Main Street",
        })
        self.assertRedirects(response, reverse("facility_detail", args=[self.facility.id]))
        self.facility.refresh_from_db()
        self.assertEqual(self.facility.hospital_name, "Updated Hospital")
        self.assertEqual(self.client.get(reverse("facility_delete", args=[self.facility.id])).status_code, 405)
        response = self.client.post(reverse("facility_delete", args=[self.facility.id]))
        self.assertRedirects(response, reverse("facility_list"))
        self.assertFalse(FacilityInformation.objects.filter(id=self.facility.id).exists())

    def test_all_seven_mou_types_generate_and_record_history(self):
        for mou_type in TEMPLATE_FILES:
            response = self.client.get(reverse("generate_mou", args=[self.facility.id, mou_type]))
            self.assertEqual(response.status_code, 200, mou_type)
            self.assertEqual(response["Content-Type"], "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
            response.close()
        self.assertEqual(MOUHistory.objects.count(), 7)

    def test_add_form_includes_saved_provider_lookup_for_name_recommendations(self):
        FacilityInformation.objects.create(
            hospital_name="City Hospital",
            hospital_address="12 Market Road",
            ambulance_name="City Ambulance",
            ambulance_address="12 Market Road",
        )

        response = self.client.get(reverse("facility_add"))

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, "provider-address-data")
        self.assertContains(response, "\"ambulance_name\"")
        self.assertContains(response, "\"City Ambulance\"")
