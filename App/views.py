import os
from datetime import timedelta
from pathlib import Path
from zipfile import BadZipFile

from django.conf import settings
from django.contrib import messages
from django.contrib.auth.decorators import login_required
from django.db.models import Q
from django.http import FileResponse, HttpResponse
from django.shortcuts import render, redirect, get_object_or_404
from django.utils import timezone

from docx import Document

from .forms import FacilityInformationForm
from .models import FacilityInformation, MOUHistory

TEMPLATE_FILES = {
    "ambulance": ("ambulance.docx", "Ambulance.docx"),
    "blood_bank": ("blood_bank.docx", "Blood Bank.docx"),
    "canteen": ("canteen.docx", "Canteen.docx"),
    "second_hospital": ("second_hospital.docx", "Emergency.docx"),
    "lab": ("lab.docx", "Lab.docx"),
    "radio_lab": ("radio_lab.docx", "Radio.docx"),
    "dry_cleaner": ("dry_cleaner.docx", "Dry Cleaners.docx"),
}

MOU_TYPES = (
    ("ambulance", "Ambulance MOU", "🚑", "Ambulance Provider", "ambulance_name", "ambulance_address"),
    ("blood_bank", "Blood Bank Agreement", "🩸", "Blood Bank", "blood_bank_name", "blood_bank_address"),
    ("canteen", "Canteen MOU", "🍽", "Canteen", "canteen_name", "canteen_address"),
    ("second_hospital", "Emergency / Second Hospital MOU", "🏥", "Second Hospital", "second_hospital_name", "second_hospital_address"),
    ("lab", "Laboratory MOU", "🧪", "Laboratory", "lab_name", "lab_address"),
    ("radio_lab", "Radio Lab MOU", "📡", "Radio Lab", "radio_lab_name", "radio_lab_address"),
    ("dry_cleaner", "Dry Cleaner MOU", "👕", "Dry Cleaner", "dry_cleaner_name", "dry_cleaner_address"),
)

PROVIDER_DETAIL_FIELDS = (
    ("dry_cleaner", "dry_cleaner_name", "dry_cleaner_address"),
    ("blood_bank", "blood_bank_name", "blood_bank_address"),
    ("canteen", "canteen_name", "canteen_address"),
    ("second_hospital", "second_hospital_name", "second_hospital_address"),
    ("ambulance", "ambulance_name", "ambulance_address"),
    ("radio_lab", "radio_lab_name", "radio_lab_address"),
    ("lab", "lab_name", "lab_address"),
)


def _provider_address_data():
    provider_data = []
    for facility in FacilityInformation.objects.all():
        for provider_name, name_field, address_field in PROVIDER_DETAIL_FIELDS:
            name = getattr(facility, name_field, "")
            address = getattr(facility, address_field, "")
            if name and address:
                provider_data.append({
                    "provider": provider_name,
                    "name": name,
                    "address": address,
                    "name_field": name_field,
                    "address_field": address_field,
                })
    return provider_data


@login_required
def facility_information(request):
    if request.method == "POST":
        form = FacilityInformationForm(request.POST)
        if form.is_valid():
            facility = form.save()
            messages.success(request, "Facility added successfully.")
            return redirect("dashboard")
    else:
        form = FacilityInformationForm()

    return render(request, "facility_information.html", {
        "form": form,
        "provider_address_data": _provider_address_data(),
    })


def _contract_status(facility, default="Generated"):
    today = timezone.localdate()
    if facility.contract_end_date and facility.contract_end_date < today:
        return "Expired"
    if facility.contract_start_date and facility.contract_end_date:
        if facility.contract_start_date <= today <= facility.contract_end_date:
            if (facility.contract_end_date - today).days <= 30:
                return "Expiring Soon"
            return "Active"
    return default


def _history_rows(queryset):
    rows = list(queryset)
    for history in rows:
        history.display_status = _contract_status(history.facility, history.status)
    return rows


@login_required
def dashboard(request):
    histories = MOUHistory.objects.select_related("facility")
    facilities = FacilityInformation.objects.all()
    month_start = timezone.localdate().replace(day=1)
    today = timezone.localdate()
    expiring = histories.filter(
        facility__contract_end_date__gte=today,
        facility__contract_end_date__lte=today + timedelta(days=30),
    )
    context = {
        "facility_count": facilities.count(),
        "facility_month_count": facilities.filter(created_at__date__gte=month_start).count(),
        "mou_count": histories.count(),
        "active_count": sum(row.display_status == "Active" for row in _history_rows(histories)),
        "expiring_count": expiring.count(),
        "recent_history": _history_rows(histories[:8]),
        "expiring_history": _history_rows(expiring.order_by("facility__contract_end_date")[:8]),
    }
    return render(request, "dashboard.html", context)


@login_required
def facility_list(request):
    query = request.GET.get("q", "").strip()
    facilities = FacilityInformation.objects.all()
    if query:
        facilities = facilities.filter(Q(hospital_name__icontains=query) | Q(hospital_address__icontains=query))
    return render(request, "facility_list.html", {"facilities": facilities, "query": query})


@login_required
def facility_add(request):
    if request.method == "POST":
        form = FacilityInformationForm(request.POST)
        if form.is_valid():
            facility = form.save()
            messages.success(request, "Facility added successfully.")
            return redirect("facility_detail", facility_id=facility.id)
    else:
        form = FacilityInformationForm()
    return render(request, "facility_edit.html", {
        "form": form,
        "page_title": "Add Facility",
        "provider_address_data": _provider_address_data(),
    })


@login_required
def facility_detail(request, facility_id):
    facility = get_object_or_404(FacilityInformation, id=facility_id)
    provider_rows = [
        ("Ambulance", facility.ambulance_name, facility.ambulance_address),
        ("Blood Bank", facility.blood_bank_name, facility.blood_bank_address),
        ("Canteen", facility.canteen_name, facility.canteen_address),
        ("Second Hospital", facility.second_hospital_name, facility.second_hospital_address),
        ("Radio Lab", facility.radio_lab_name, facility.radio_lab_address),
        ("Laboratory", facility.lab_name, facility.lab_address),
        ("Dry Cleaner", facility.dry_cleaner_name, facility.dry_cleaner_address),
    ]
    return render(request, "facility_detail.html", {
        "facility": facility,
        "history": _history_rows(facility.mou_history.all()[:8]),
        "provider_rows": provider_rows,
    })


@login_required
def facility_edit(request, facility_id):
    facility = get_object_or_404(FacilityInformation, id=facility_id)
    if request.method == "POST":
        form = FacilityInformationForm(request.POST, instance=facility)
        if form.is_valid():
            form.save()
            messages.success(request, "Facility information updated successfully.")
            return redirect("facility_detail", facility_id=facility.id)
    else:
        form = FacilityInformationForm(instance=facility)
    return render(request, "facility_edit.html", {
        "form": form,
        "facility": facility,
        "page_title": "Edit Facility",
        "provider_address_data": _provider_address_data(),
    })


@login_required
def facility_delete(request, facility_id):
    if request.method != "POST":
        return HttpResponse("Delete requires a POST request.", status=405)
    facility = get_object_or_404(FacilityInformation, id=facility_id)
    facility.delete()
    messages.success(request, "Facility deleted successfully.")
    return redirect("facility_list")


@login_required
def mou_documents(request):
    return render(request, "mou_documents.html", {"mou_types": MOU_TYPES, "facilities": FacilityInformation.objects.all()})


@login_required
def generate_mou_page(request):
    facilities = FacilityInformation.objects.all()
    selected_facility = None
    selected_type = request.GET.get("mou_type", "ambulance")
    if request.method == "POST":
        selected_facility = get_object_or_404(FacilityInformation, id=request.POST.get("facility_id"))
        selected_type = request.POST.get("mou_type", "")
        return _generate_mou_response(request, selected_facility, selected_type)
    facility_id = request.GET.get("facility_id")
    if facility_id:
        selected_facility = get_object_or_404(FacilityInformation, id=facility_id)
    return render(request, "generate_mou.html", {
        "facilities": facilities,
        "selected_facility": selected_facility,
        "selected_type": selected_type,
        "mou_types": MOU_TYPES,
    })


@login_required
def mou_history(request):
    query = request.GET.get("q", "").strip()
    status = request.GET.get("status", "").strip()
    histories = MOUHistory.objects.select_related("facility")
    if query:
        histories = histories.filter(Q(facility__hospital_name__icontains=query) | Q(mou_type__icontains=query))
    if status == "Active":
        histories = histories.filter(
            facility__contract_start_date__lte=timezone.localdate(),
            facility__contract_end_date__gte=timezone.localdate(),
        )
    elif status == "Expired":
        histories = histories.filter(facility__contract_end_date__lt=timezone.localdate())
    elif status in {choice[0] for choice in MOUHistory.STATUS_CHOICES}:
        histories = histories.filter(status=status)
    return render(request, "mou_history.html", {
        "history": _history_rows(histories),
        "query": query,
        "selected_status": status,
        "status_choices": MOUHistory.STATUS_CHOICES,
    })


@login_required
def mou_list(request, facility_id):
    facility = get_object_or_404(FacilityInformation, id=facility_id)
    return render(request, "mou_list.html", {"facility": facility})


def _build_replacements(facility):
    field_values = {
        "hospital_name": facility.hospital_name,
        "hospital_address": facility.hospital_address,
        "dry_cleaner_name": facility.dry_cleaner_name,
        "dry_cleaner_address": facility.dry_cleaner_address,
        "blood_bank_name": facility.blood_bank_name,
        "blood_bank_address": facility.blood_bank_address,
        "canteen_name": facility.canteen_name,
        "canteen_address": facility.canteen_address,
        "second_hospital_name": facility.second_hospital_name,
        "second_hospital_address": facility.second_hospital_address,
        "ambulance_name": facility.ambulance_name,
        "ambulance_address": facility.ambulance_address,
        "radio_lab_name": facility.radio_lab_name,
        "radio_lab_address": facility.radio_lab_address,
        "lab_name": facility.lab_name,
        "lab_address": facility.lab_address,
    }
    replacements = {}
    for field_name, value in field_values.items():
        replacements[f"{{{field_name}}}"] = value or ""
        replacements[f"{{{field_name.upper()}}}"] = value or ""

    replacements.update({
        "{EMERGENCY_NAME}": facility.second_hospital_name or "",
        "{EMERGENCY_ADDRESS}": facility.second_hospital_address or "",
        "{RADIO_NAME}": facility.radio_lab_name or "",
        "{RADIO_ADDRESS}": facility.radio_lab_address or "",
        "{DRY_NAME}": facility.dry_cleaner_name or "",
        "{DRY_ADDRESS}": facility.dry_cleaner_address or "",
        "2nd party hospital name": facility.second_hospital_name or "",
        "Near Park Road, 9-Civil Lines Gorakhpur Uttar Pradesh": facility.second_hospital_address or "",
        "Get Well Hospital and Ultrasound Pvt Ltd.": facility.hospital_name or "",
        "Rishab Singh catering services": facility.canteen_name or "",
        "Mr. Rishab Singh": facility.canteen_name or "",
        "Salekh Chand Catering Services": facility.canteen_name or "",
        "Arogya BLOOD BANK & COMPONENT CENTRE CHARITABLE": facility.blood_bank_name or "",
        "GORAKHPUR BLOOD BANK": facility.blood_bank_name or "",
        "Gorakhpur BLOOD BANK": facility.blood_bank_name or "",
        "Health Care Imaging Centre": facility.radio_lab_name or "",
        "Health Care Imaging Center": facility.radio_lab_name or "",
        "Anjaneya Superspeciality Hospital and Neuro Centre": facility.hospital_name or "",
        "Saini Orthopedic Super Specialty Centre": facility.hospital_name or "",
        "Raptinagar phase 1 Gorakhpur 273003": facility.hospital_address or "",
        "HOSPITAL NAME": facility.hospital_name or "",
    })
    return replacements


def replace_text_in_paragraph(paragraph, replacements):
    while paragraph.runs:
        full_text = "".join(run.text or "" for run in paragraph.runs)
        matches = [
            (full_text.find(placeholder), placeholder, str(value))
            for placeholder, value in replacements.items()
            if full_text.find(placeholder) >= 0
        ]
        if not matches:
            return

        start, placeholder, value = min(matches, key=lambda match: match[0])
        end = start + len(placeholder)
        positions = []
        cursor = 0
        for run in paragraph.runs:
            run_start = cursor
            cursor += len(run.text or "")
            positions.append((run, run_start, cursor))

        start_run, start_offset = next(
            (run, start - run_start)
            for run, run_start, run_end in positions
            if run_start <= start < run_end or (start == run_end == run_start and not run.text)
        )
        end_run, end_offset = next(
            (run, end - run_start)
            for run, run_start, run_end in positions
            if run_start < end <= run_end
        )

        if start_run is end_run:
            start_run.text = start_run.text[:start_offset] + value + start_run.text[end_offset:]
            continue

        start_run.text = start_run.text[:start_offset] + value
        clearing = False
        for run, _, _ in positions:
            if run is start_run:
                clearing = True
                continue
            if clearing and run is not end_run:
                run.text = ""
            if run is end_run:
                run.text = run.text[end_offset:]
                break


def replace_text_in_table(table, replacements):
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                replace_text_in_paragraph(paragraph, replacements)
            for nested_table in cell.tables:
                replace_text_in_table(nested_table, replacements)


def replace_text_in_document(document, replacements):
    for paragraph in document.paragraphs:
        replace_text_in_paragraph(paragraph, replacements)

    for table in document.tables:
        replace_text_in_table(table, replacements)

    for section in document.sections:
        for part in (section.header, section.footer):
            for paragraph in part.paragraphs:
                replace_text_in_paragraph(paragraph, replacements)
            for table in part.tables:
                replace_text_in_table(table, replacements)


def _generate_mou_response(request, facility, mou_type):
    normalized_type = mou_type.strip().lower()
    template_names = TEMPLATE_FILES.get(normalized_type)

    if not template_names:
        return HttpResponse("Invalid MOU type.", status=400)

    template_directory = Path(settings.BASE_DIR) / "App" / "mou_templates"
    template_path = next(
        (template_directory / name for name in template_names if (template_directory / name).is_file()),
        None,
    )

    if template_path is None:
        return HttpResponse("The selected MOU template is unavailable.", status=404)

    try:
        document = Document(template_path)
    except (BadZipFile, OSError, ValueError):
        return HttpResponse("The selected MOU template could not be opened.", status=500)
    replacements = _build_replacements(facility)
    replace_text_in_document(document, replacements)

    output_directory = Path(settings.BASE_DIR) / "generated_mous"
    os.makedirs(output_directory, exist_ok=True)

    output_file = output_directory / f"{normalized_type}_MOU_{facility.id}.docx"
    try:
        document.save(output_file)
    except OSError:
        return HttpResponse("The generated MOU could not be saved.", status=500)

    MOUHistory.objects.create(
        facility=facility,
        mou_type=dict((item[0], item[1]) for item in MOU_TYPES).get(normalized_type, normalized_type),
        file_name=output_file.name,
    )
    messages.success(request, "MOU generated successfully.")
    return FileResponse(
        output_file.open("rb"),
        as_attachment=True,
        filename=f"{normalized_type}_MOU.docx",
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


@login_required
def generate_mou(request, facility_id, mou_type):
    facility = get_object_or_404(FacilityInformation, id=facility_id)
    return _generate_mou_response(request, facility, mou_type)


@login_required
def download_history(request, history_id):
    history = get_object_or_404(MOUHistory.objects.select_related("facility"), id=history_id)
    output_directory = Path(settings.BASE_DIR) / "generated_mous"
    file_path = output_directory / Path(history.file_name).name
    if not history.file_name or not file_path.is_file():
        return _generate_mou_response(request, history.facility, _type_key(history.mou_type))
    return FileResponse(
        file_path.open("rb"),
        as_attachment=True,
        filename=file_path.name,
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


def _type_key(label):
    for key, mou_label, *_ in MOU_TYPES:
        if label == key or label == mou_label:
            return key
    return label